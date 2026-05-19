from __future__ import annotations

import asyncio
import json
import os
import shutil
import time
import uuid
from concurrent.futures import ThreadPoolExecutor
from io import BytesIO
from pathlib import Path
from typing import Any, Literal, Optional

import orjson
from fastapi import Depends, FastAPI, File, Form, Header, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, PlainTextResponse, RedirectResponse, Response, StreamingResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from pydantic import BaseModel, Field

from .billing import BillingService
from .douyin import is_douyin_url
from .jobs import JobStore, ensure_job_dir
from .models import JobProgress, JobStatus, SummaryStatus
from .settings import settings
from .summary import summarize_with_deepseek
from .transcript import diagnose_subtitles, extract_transcript, is_bilibili_url
from .ydl import download_url


app = FastAPI(title="Free Video Downloader", version="0.1.0")
BASE_DIR = Path(__file__).resolve().parent

templates = Jinja2Templates(directory=str(BASE_DIR / "templates"))
app.mount("/static", StaticFiles(directory=str(BASE_DIR / "static")), name="static")

job_store = JobStore()
billing_service = BillingService()
executor = ThreadPoolExecutor(max_workers=6)
summary_semaphore = asyncio.Semaphore(settings.max_summary_workers)

DEVICE_COOKIE = "fvd_uid"


def _client_ip(req: Request) -> str:
    # Minimal; behind proxy should be configured with trusted headers by deploy layer.
    return req.client.host if req.client else "unknown"


def _device_id(request: Request) -> str:
    return getattr(request.state, "device_id", "")


def _require_admin(x_admin_token: str | None) -> None:
    if not settings.admin_token:
        raise HTTPException(status_code=403, detail="Admin token not configured")
    if not x_admin_token or x_admin_token != settings.admin_token:
        raise HTTPException(status_code=403, detail="Invalid admin token")


@app.middleware("http")
async def _device_cookie_middleware(request: Request, call_next):
    device_id = request.cookies.get(DEVICE_COOKIE)
    set_cookie = False
    if not device_id:
        device_id = uuid.uuid4().hex
        set_cookie = True
    request.state.device_id = device_id
    response = await call_next(request)
    if set_cookie:
        response.set_cookie(
            DEVICE_COOKIE,
            device_id,
            httponly=True,
            samesite="lax",
            max_age=365 * 24 * 3600,
        )
    return response


@app.on_event("startup")
async def _startup() -> None:
    if settings.enable_ai_summary and (not settings.demo_mode) and (not settings.deepseek_api_key.strip()):
        raise RuntimeError(
            "FVD_DEEPSEEK_API_KEY is required when AI summary is enabled. "
            "Please inject it via environment variable."
        )

    billing_service.init()
    app.state.loop = asyncio.get_running_loop()

    async def cleaner() -> None:
        while True:
            try:
                await job_store.cleanup_expired()
            finally:
                await asyncio.sleep(settings.cleanup_interval_seconds)

    asyncio.create_task(cleaner())


@app.get("/", response_class=HTMLResponse)
async def index(request: Request) -> Any:
    billing = billing_service.get_billing_status(_device_id(request))
    return templates.TemplateResponse(
        "index.html",
        {"request": request, "title": "万能视频下载 · 一键保存", "billing": billing},
    )


class CheckoutBody(BaseModel):
    plan: Literal["monthly", "yearly"] = Field(default="monthly")
    idempotency_key: str | None = None


@app.get("/api/billing/me")
async def billing_me(request: Request) -> Any:
    return billing_service.get_billing_status(_device_id(request))


@app.post("/api/billing/checkout-session")
async def create_checkout_session(request: Request, body: CheckoutBody) -> Any:
    device_id = _device_id(request)
    try:
        result = billing_service.create_checkout_session(
            device_id=device_id,
            plan=body.plan,
            idempotency_key=body.idempotency_key,
        )
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e
    return result


@app.post("/api/stripe/webhook")
async def stripe_webhook(request: Request) -> Response:
    payload = await request.body()
    sig = request.headers.get("Stripe-Signature")
    try:
        billing_service.handle_webhook(payload, sig)
    except ValueError as e:
        raise HTTPException(status_code=400, detail=str(e)) from e
    except RuntimeError as e:
        raise HTTPException(status_code=503, detail=str(e)) from e
    return Response(status_code=200)


@app.get("/billing/success", response_class=HTMLResponse)
async def billing_success(request: Request, session_id: str | None = None) -> Any:
    device_id = _device_id(request)
    verified: dict[str, Any] = {"ok": False}
    if session_id:
        try:
            verified = billing_service.verify_checkout_session(session_id, device_id)
        except Exception:
            verified = {"ok": billing_service.is_pro(device_id), "session_id": session_id}
    billing = billing_service.get_billing_status(device_id)
    return templates.TemplateResponse(
        "billing_success.html",
        {
            "request": request,
            "title": "开通成功 · Pro",
            "billing": billing,
            "session_id": session_id,
            "verified": verified,
        },
    )


@app.get("/billing/cancel", response_class=HTMLResponse)
async def billing_cancel(request: Request) -> Any:
    return templates.TemplateResponse(
        "billing_cancel.html",
        {"request": request, "title": "已取消支付"},
    )


@app.get("/billing/demo-checkout", response_class=HTMLResponse)
async def billing_demo_checkout(request: Request, plan: str = "monthly") -> Any:
    if not settings.stripe_demo_mode:
        raise HTTPException(status_code=404, detail="Not found")
    device_id = _device_id(request)
    normalized: Literal["monthly", "yearly"] = "yearly" if plan == "yearly" else "monthly"
    billing_service.activate_demo_subscription(device_id, plan=normalized)
    return RedirectResponse(url="/billing/success?session_id=demo_ok", status_code=303)


def _parse_urls(raw: str) -> list[str]:
    urls: list[str] = []
    for part in raw.replace("\r", "\n").split("\n"):
        s = part.strip()
        if not s:
            continue
        if len(s) > 2000:
            continue
        if s.startswith("http://") or s.startswith("https://"):
            # 抖音 URL 验证 - 只接受视频链接
            if is_douyin_url(s):
                # 抖音链接由 DouyinParser 处理（无需 Cookie）
                urls.append(s)
            elif "tiktok.com" in s:
                # TikTok 链接检查格式
                if "/video/" in s or "/v/" in s or "v.tiktok.com" in s:
                    urls.append(s)
            else:
                # 其他平台链接
                urls.append(s)
    return urls


def _human_error(e: Exception) -> str:
    s = str(e) or e.__class__.__name__
    
    # 针对抖音 URL 的特殊处理
    if "Unsupported URL" in s and "douyin.com" in s:
        return "不支持的抖音链接格式。请使用视频链接（如 https://v.douyin.com/... 或 https://www.douyin.com/video/...）"
    
    # Remove overly long noisy prefixes
    return s[:1200]


def _normalize_summary_template(raw: str | None) -> str:
    allowed = {"learning", "course"}
    v = (raw or "learning").strip().lower()
    return v if v in allowed else "learning"


def _summary_to_markdown(job: Any) -> str:
    s = job.summary_result or {}
    lines: list[str] = []
    lines.append(f"# 视频总结：{job.display_name or job.url}")
    lines.append("")
    lines.append(f"- 任务ID：`{job.id}`")
    lines.append(f"- 来源链接：{job.url}")
    lines.append(f"- 模板：{(s.get('template') or job.meta.get('summary_template') or 'learning')}")
    lines.append(f"- 模型：{s.get('model', 'unknown')}")
    lines.append("")

    lines.append("## 视频大纲")
    for item in s.get("outline", []):
        lines.append(f"- {item}")
    lines.append("")

    lines.append("## 核心要点")
    for item in s.get("key_points", []):
        lines.append(f"- {item}")
    lines.append("")

    lines.append("## 行动建议")
    for item in s.get("action_items", []):
        lines.append(f"- {item}")
    lines.append("")

    lines.append("## 关键词术语")
    for item in s.get("terms", []):
        lines.append(f"- {item}")
    lines.append("")
    return "\n".join(lines).strip() + "\n"


def _transcript_to_text(transcript: dict[str, Any]) -> str:
    lines: list[str] = []
    lines.append("# 字幕/转写文本")
    lines.append("")
    lines.append(f"- 来源：{transcript.get('source', 'unknown')}")
    lines.append(f"- 语言：{transcript.get('lang', 'unknown')}")
    if transcript.get("title"):
        lines.append(f"- 标题：{transcript.get('title')}")
    lines.append("")
    lines.append("## 正文")
    lines.append(transcript.get("text", ""))
    return "\n".join(lines).strip() + "\n"


def _summary_to_docx_bytes(job: Any) -> bytes:
    try:
        from docx import Document  # type: ignore
    except Exception as e:
        raise RuntimeError(f"未安装 python-docx，无法导出 Word: {e}")

    s = job.summary_result or {}
    doc = Document()
    doc.add_heading(f"视频总结：{job.display_name or job.url}", level=1)
    doc.add_paragraph(f"任务ID：{job.id}")
    doc.add_paragraph(f"来源链接：{job.url}")
    doc.add_paragraph(f"模板：{(s.get('template') or job.meta.get('summary_template') or 'learning')}")
    doc.add_paragraph(f"模型：{s.get('model', 'unknown')}")

    doc.add_heading("视频大纲", level=2)
    for item in s.get("outline", []):
        doc.add_paragraph(str(item), style="List Bullet")

    doc.add_heading("核心要点", level=2)
    for item in s.get("key_points", []):
        doc.add_paragraph(str(item), style="List Bullet")

    doc.add_heading("行动建议", level=2)
    for item in s.get("action_items", []):
        doc.add_paragraph(str(item), style="List Bullet")

    doc.add_heading("关键词术语", level=2)
    for item in s.get("terms", []):
        doc.add_paragraph(str(item), style="List Bullet")

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _job_to_dict(j) -> dict[str, Any]:
    transcript_path = Path(settings.download_root) / j.id / "transcript.json"
    return {
        "id": j.id,
        "url": j.url,
        "status": j.status.value,
        "progress": {
            "status": j.progress.status.value,
            "percent": j.progress.percent,
            "speed": j.progress.speed,
            "eta": j.progress.eta,
            "filename": j.progress.filename,
            "message": j.progress.message,
        },
        "display_name": j.display_name,
        "error": j.error,
        "expires_at": j.expires_at,
        "summary_status": j.summary_status.value,
        "summary_error": j.summary_error,
        "transcript_error": j.transcript_error,
        "summary_result": j.summary_result,
        "transcript_available": transcript_path.exists(),
        "meta": j.meta,
    }


async def _run_summary_pipeline(job_id: str, url: str, output_path: str, job_dir: str) -> None:
    try:
        job = await job_store.get_job(job_id)
        if not job.meta.get("enable_summary", False):
            await job_store.set_summary_status(job_id, SummaryStatus.skipped)
            return

        async with summary_semaphore:
            # 第一步：提取字幕
            await job_store.set_summary_status(job_id, SummaryStatus.transcribing)
            transcript = None
            
            try:
                transcript = await asyncio.to_thread(
                    extract_transcript,
                    url=url,
                    job_dir=job_dir,
                    output_path=output_path,
                )
                transcript_path = Path(job_dir) / "transcript.json"
                transcript_path.write_text(orjson.dumps(transcript).decode("utf-8"), encoding="utf-8")
            except Exception as e:
                # 字幕提取失败，记录错误信息但不直接中断
                transcript_error_msg = _human_error(e)
                await job_store.set_summary_status(
                    job_id,
                    SummaryStatus.failed,
                    transcript_error=transcript_error_msg,
                    summary_error="字幕提取失败，无法生成视频总结。但视频文件已下载完成。",
                )
                return

            # 检查是否有字幕
            if transcript.get("source") == "none" or not transcript.get("text", "").strip():
                await job_store.set_summary_status(
                    job_id,
                    SummaryStatus.failed,
                    transcript_error=transcript.get("error", "该视频没有字幕"),
                    summary_error="视频无字幕，无法生成总结。但视频文件已下载完成。",
                )
                return

            # 第二步：生成总结
            await job_store.set_summary_status(job_id, SummaryStatus.summarizing)
            try:
                summary_result = await asyncio.to_thread(
                    summarize_with_deepseek,
                    transcript_text=transcript.get("text", ""),
                    video_title=transcript.get("title"),
                    video_url=url,
                    summary_template=job.meta.get("summary_template", "learning"),
                )
            except Exception as e:
                await job_store.set_summary_status(
                    job_id,
                    SummaryStatus.failed,
                    summary_error=_human_error(e),
                )
                return

            summary_result["transcript_source"] = transcript.get("source")
            summary_result["transcript_lang"] = transcript.get("lang")
            summary_path = Path(job_dir) / "summary.json"
            summary_path.write_text(orjson.dumps(summary_result).decode("utf-8"), encoding="utf-8")
            await job_store.set_summary_status(job_id, SummaryStatus.done, summary_result=summary_result)
    except Exception as e:
        await job_store.set_summary_status(job_id, SummaryStatus.failed, summary_error=_human_error(e))


@app.post("/api/jobs")
async def create_jobs(
    request: Request,
    urls: str = Form(...),
    summary_template: str = Form(default="learning"),
    cookies: Optional[UploadFile] = File(default=None),
    x_admin_token: Optional[str] = Header(default=None, alias="X-Admin-Token"),
) -> Any:
    ip = _client_ip(request)
    url_list = _parse_urls(urls)
    if not url_list:
        raise HTTPException(status_code=400, detail="No valid URLs provided")
    if len(url_list) > settings.max_urls_per_request:
        raise HTTPException(status_code=400, detail=f"Too many URLs (max {settings.max_urls_per_request})")

    ent = billing_service.get_entitlements(_device_id(request))
    max_active = ent["max_active_jobs_per_ip"]

    active = await job_store.count_active_for_ip(ip)
    if active >= max_active:
        raise HTTPException(status_code=429, detail="Too many active downloads for this IP")

    cookies_path: Optional[str] = None
    if cookies is not None:
        if not settings.enable_cookies:
            raise HTTPException(status_code=403, detail="cookies is disabled by default")
        _require_admin(x_admin_token)

    jobs = []
    summary_template = _normalize_summary_template(summary_template)
    for u in url_list:
        enable_summary = settings.enable_ai_summary and (
            ent["summary_all_platforms"] or is_bilibili_url(u)
        )
        initial_summary_status = SummaryStatus.queued if enable_summary else SummaryStatus.skipped
        job = await job_store.create_job(
            url=u,
            ip=ip,
            meta={"enable_summary": enable_summary, "summary_template": summary_template},
            max_active_jobs_per_ip=max_active,
        )
        await job_store.update_job(job.id, summary_status=initial_summary_status)
        jobs.append(job)

    # If cookies uploaded, save once and reuse for this batch (same request).
    if cookies is not None:
        job_dir = ensure_job_dir(jobs[0].id)
        cookies_path = str(Path(job_dir) / "cookies.txt")
        content = await cookies.read()
        Path(cookies_path).write_bytes(content)

    loop = getattr(app.state, "loop", asyncio.get_running_loop())

    async def run_job(job_id: str, url: str, cookies_path_for_job: Optional[str]) -> None:
        job_dir = ensure_job_dir(job_id)
        if cookies_path_for_job and cookies_path_for_job != str(Path(job_dir) / "cookies.txt"):
            # copy cookies into this job dir if shared file
            try:
                Path(job_dir).mkdir(parents=True, exist_ok=True)
                shutil.copyfile(cookies_path_for_job, str(Path(job_dir) / "cookies.txt"))
            except Exception:
                pass
            cookies_path_for_job = str(Path(job_dir) / "cookies.txt")

        def on_progress(d: dict[str, Any]) -> None:
            status = d.get("status")
            if loop.is_closed():
                return
            if status == "downloading":
                percent_str = d.get("_percent_str") or ""
                try:
                    percent = float(percent_str.replace("%", "").strip())
                except Exception:
                    percent = None
                prog = JobProgress(
                    status=JobStatus.downloading,
                    percent=percent,
                    speed=d.get("_speed_str"),
                    eta=d.get("_eta_str"),
                    filename=os.path.basename(d.get("filename", "")) if d.get("filename") else None,
                    message=None,
                )
                asyncio.run_coroutine_threadsafe(job_store.set_progress(job_id, prog), loop)
            elif status == "finished":
                prog = JobProgress(status=JobStatus.downloading, percent=100.0, message="Processing...")
                asyncio.run_coroutine_threadsafe(job_store.set_progress(job_id, prog), loop)

        try:
            await job_store.set_progress(job_id, JobProgress(status=JobStatus.downloading, message="Starting..."))
            output_path, display_name = await loop.run_in_executor(
                executor,
                lambda: download_url(url=url, job_dir=job_dir, on_progress=on_progress, cookies_path=cookies_path_for_job),
            )
            await job_store.mark_finished(job_id, output_path=output_path, display_name=display_name)
            
            # 总结流程单独处理，确保其失败不会影响已下载完成的状态
            try:
                await _run_summary_pipeline(job_id, url, output_path, job_dir)
            except Exception as summary_err:
                # 总结流程失败不影响下载完成的状态，只记录总结错误
                pass
                
        except Exception as e:
            await job_store.mark_failed(job_id, error=_human_error(e))

    # fire and forget each job
    for j in jobs:
        asyncio.create_task(run_job(j.id, j.url, cookies_path))

    return Response(
        content=orjson.dumps({"jobs": [_job_to_dict(j) for j in jobs]}),
        media_type="application/json",
    )


@app.get("/api/jobs/{job_id}")
async def get_job(job_id: str) -> Any:
    j = await job_store.get_job(job_id)
    return _job_to_dict(j)


@app.get("/api/jobs/{job_id}/events")
async def job_events(job_id: str) -> StreamingResponse:
    async def gen():
        # Initial ping to establish connection
        yield "event: ping\ndata: {}\n\n"
        async for evt in job_store.subscribe(job_id):
            data = orjson.dumps(evt).decode("utf-8")
            yield f"data: {data}\n\n"

    return StreamingResponse(gen(), media_type="text/event-stream")


@app.get("/api/jobs/{job_id}/file")
async def download_file(job_id: str) -> FileResponse:
    j = await job_store.get_job(job_id)
    if j.status != JobStatus.finished or not j.output_path:
        raise HTTPException(status_code=409, detail="Job not finished")
    p = Path(j.output_path)
    if not p.exists():
        raise HTTPException(status_code=404, detail="File not found")
    filename = j.display_name or p.name
    return FileResponse(path=str(p), filename=filename, media_type="application/octet-stream")


@app.get("/api/jobs/{job_id}/summary")
async def get_summary(job_id: str) -> Any:
    j = await job_store.get_job(job_id)
    if j.summary_status != SummaryStatus.done:
        raise HTTPException(status_code=409, detail=f"Summary not ready: {j.summary_status.value}")
    return {
        "id": j.id,
        "summary_status": j.summary_status.value,
        "summary_result": j.summary_result,
    }


@app.get("/api/jobs/{job_id}/transcript")
async def get_transcript(job_id: str) -> Any:
    await job_store.get_job(job_id)
    transcript_path = Path(settings.download_root) / job_id / "transcript.json"
    if not transcript_path.exists():
        raise HTTPException(status_code=409, detail="Transcript not ready")
    try:
        transcript = orjson.loads(transcript_path.read_bytes())
    except Exception:
        raise HTTPException(status_code=500, detail="Transcript file is corrupted")
    return {"id": job_id, "transcript": transcript}


@app.post("/api/jobs/{job_id}/summary/retry")
async def retry_summary(job_id: str) -> Any:
    j = await job_store.get_job(job_id)
    if j.status != JobStatus.finished or not j.output_path:
        raise HTTPException(status_code=409, detail="Job not finished")
    if not j.meta.get("enable_summary", False):
        raise HTTPException(status_code=400, detail="Summary is disabled for this job")

    await job_store.set_summary_status(job_id, SummaryStatus.queued, summary_error=None, summary_result=None)
    job_dir = ensure_job_dir(job_id)
    asyncio.create_task(_run_summary_pipeline(job_id, j.url, j.output_path, job_dir))
    return {"id": j.id, "summary_status": SummaryStatus.queued.value}


@app.get("/api/jobs/{job_id}/summary.md")
async def export_summary_markdown(job_id: str) -> PlainTextResponse:
    j = await job_store.get_job(job_id)
    if j.summary_status != SummaryStatus.done:
        raise HTTPException(status_code=409, detail=f"Summary not ready: {j.summary_status.value}")
    content = _summary_to_markdown(j)
    headers = {"Content-Disposition": f'attachment; filename="summary-{job_id}.md"'}
    return PlainTextResponse(content=content, media_type="text/markdown; charset=utf-8", headers=headers)


@app.get("/api/jobs/{job_id}/transcript.txt")
async def export_transcript_text(job_id: str) -> PlainTextResponse:
    await job_store.get_job(job_id)
    transcript_path = Path(settings.download_root) / job_id / "transcript.json"
    if not transcript_path.exists():
        raise HTTPException(status_code=409, detail="Transcript not ready")
    try:
        transcript = orjson.loads(transcript_path.read_bytes())
    except Exception:
        raise HTTPException(status_code=500, detail="Transcript file is corrupted")
    content = _transcript_to_text(transcript)
    headers = {"Content-Disposition": f'attachment; filename="transcript-{job_id}.txt"'}
    return PlainTextResponse(content=content, media_type="text/plain; charset=utf-8", headers=headers)


@app.get("/api/jobs/{job_id}/summary.docx")
async def export_summary_docx(job_id: str) -> Response:
    j = await job_store.get_job(job_id)
    if j.summary_status != SummaryStatus.done:
        raise HTTPException(status_code=409, detail=f"Summary not ready: {j.summary_status.value}")
    data = _summary_to_docx_bytes(j)
    headers = {"Content-Disposition": f'attachment; filename="summary-{job_id}.docx"'}
    return Response(
        content=data,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers=headers,
    )


@app.get("/api/diagnose-subtitles")
async def api_diagnose_subtitles(url: str) -> Any:
    """诊断视频字幕信息，用于调试"""
    if not url:
        raise HTTPException(status_code=400, detail="缺少 url 参数")
    try:
        result = await asyncio.to_thread(diagnose_subtitles, url)
        return result
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"诊断失败: {str(e)}")
